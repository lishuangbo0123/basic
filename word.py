
import re
import docx
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

#水平居中
def set_tab_center(paragraph):
    # doc.tables[0].cell(a, b).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    paragraph.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER #水平居中


pattern3 = re.compile(r"^\d{1,2}\.\d")   #3级标题正则
pattern2 = re.compile(r"^\d{1,2}[^\d]")   #2级标题正则

mark_url = input('请输入文件所在的文件夹地址:')
file_list = os.listdir(mark_url)  # 待修改文件夹
# print("修改前：\n" + str(file_list))  # 输出文件夹中包含的文件
for doc_name in file_list:
    file = docx.Document(mark_url+ "\\" + doc_name)
    for paragraph in file.paragraphs:
        list = pattern3.findall(paragraph.text)
        list1 = pattern2.findall(paragraph.text)
        if paragraph.style.name == "Heading 1": #大标题，黑体三号，行距23磅
            set_tab_center(paragraph)
            for run in paragraph.runs:
                run.font.name = u'黑体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(16)
            paragraph.paragraph_format.line_spacing = Pt(23)
            # print("大标题------------------------------------------------------------------------------------" + paragraph.text)
        elif len(list) > 0:#3级标题，楷体小四，首行缩进2字符(27磅)，段前0.5行，行距23磅

            paragraph.paragraph_format.space_before = Pt(0.5)
            paragraph.paragraph_format.space_after = Pt(0.5)
            paragraph.paragraph_format.line_spacing = Pt(23)
            paragraph.style.font.size = Pt(12)
            paragraph.paragraph_format.first_line_indent = 304800
            for run in paragraph.runs:
                run.font.name = u'楷体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                run.font.size = Pt(12)
            # print("3级标题-------------------------" + paragraph.text)

        elif len(list1) > 0:#2级标题，楷体小四，首行缩进2字符(27磅)，段前0.5行，行距23磅

            paragraph.paragraph_format.space_before = Pt(0.5)
            paragraph.paragraph_format.space_after = Pt(0.5)
            paragraph.paragraph_format.line_spacing = Pt(23)
            paragraph.style.font.size = Pt(12)
            paragraph.paragraph_format.first_line_indent = 304800
            for run in paragraph.runs:
                run.font.name = u'楷体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
                run.font.size = Pt(12)
            # print("2级标题-------------------------------------------------" + paragraph.text)
        else:       #正文，宋体小四，首行缩进2字符(27磅)，行距23磅
            paragraph.paragraph_format.line_spacing = Pt(23)
            paragraph.style.font.size = Pt(12)
            paragraph.paragraph_format.first_line_indent = 304800
            # print("正文" + paragraph.text)
            for run in paragraph.runs:
                run.font.name = u'宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)

    file.save(mark_url+ "\\" + doc_name)


print('修改完成')









